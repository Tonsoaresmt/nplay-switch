from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\NplaySwitch")
OUT = ROOT / "docs" / "Nplay_Switch_Fluxo_Reproducao_e_Plano_de_Investigacao.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17223B"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PURPLE = "6D5DFB"
INK = "17202A"
MUTED = "5F6B7A"
LIGHT = "F2F4F7"
BLUE_LIGHT = "E8EEF5"
PURPLE_LIGHT = "F1EEFF"
RED = "9B1C1C"
RED_LIGHT = "FCE8E6"
GOLD = "7A5A00"
GOLD_LIGHT = "FFF4CE"
GREEN = "246B45"
GREEN_LIGHT = "E7F4EC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(p, before=0, after=6, line=1.25, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if keep:
        pf.keep_with_next = True


def add_body(doc, text="", bold_prefix=None, color=INK, after=6):
    p = doc.add_paragraph()
    set_para(p, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True, color=color)
        set_run(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        set_run(p.add_run(text), color=color)
    return p


def add_bullet(doc, text, level=0, color=INK):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_para(p, after=4, line=1.25)
    set_run(p.add_run(text), color=color)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para(p, after=5, line=1.25)
    set_run(p.add_run(text))
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    set_para(p, before=2, after=6, line=1.1)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7FA")
    p_pr.append(shd)
    set_run(p.add_run(text), size=9, color=DARK_BLUE, font="Consolas")
    return p


def add_callout(doc, label, text, fill=BLUE_LIGHT, accent=BLUE):
    p = doc.add_paragraph()
    set_para(p, before=6, after=9, line=1.2)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.12)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    set_run(p.add_run(label + "  "), bold=True, color=accent)
    set_run(p.add_run(text), color=INK)
    return p


def add_table(doc, headers, rows, widths, header_fill=BLUE_LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        set_para(p, after=0, line=1.1)
        set_run(p.add_run(header), size=9.5, bold=True, color=NAVY)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2:
                set_cell_shading(cells[i], "FAFBFC")
            p = cells[i].paragraphs[0]
            set_para(p, after=0, line=1.12)
            set_run(p.add_run(str(value)), size=9.25, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
            color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Bullet 2", "List Number"):
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para(header, after=0, line=1.0)
set_run(header.add_run("NPLAY SWITCH  |  DIAGNÓSTICO DE REPRODUÇÃO"), size=8.5, bold=True, color=MUTED)
footer = sec.footer.paragraphs[0]
add_page_number(footer)

# Capa técnica no padrão memo_masthead.
p = doc.add_paragraph()
set_para(p, before=20, after=3, line=1.0)
set_run(p.add_run("RELATÓRIO TÉCNICO"), size=10, bold=True, color=PURPLE)
p = doc.add_paragraph()
set_para(p, after=7, line=1.0)
set_run(p.add_run("Fluxo de catálogo e reprodução"), size=27, bold=True, color=NAVY)
p = doc.add_paragraph()
set_para(p, after=18, line=1.1)
set_run(p.add_run("Filmes, Séries, Animes e Doramas no Nintendo Switch"), size=15, color=DARK_BLUE)

meta = [
    ("Objetivo", "Localizar atrasos, HTTP 502 e o fechamento nativo ao iniciar HLS"),
    ("Versão analisada", "Nplay Switch 0.9.5 - commit 9ca16fd"),
    ("Arquitetura", "SDL2 + FFmpeg + libcurl + cJSON; backend Fastify/SQLite"),
    ("Data da análise", "31 de agosto de 2026"),
    ("Escopo", "Leitura do código atual; sem afirmar correção sem teste no Switch real"),
]
for label, value in meta:
    p = doc.add_paragraph()
    set_para(p, after=3, line=1.1)
    set_run(p.add_run(label + ": "), bold=True, color=NAVY)
    set_run(p.add_run(value), color=INK)

add_callout(doc, "Conclusão central", "Há pelo menos três problemas independentes: demora de catálogo/detalhes, HTTP 502 durante resolução de fonte e fechamento nativo dentro do caminho HLS. Corrigir somente um deles não torna a reprodução funcional.", PURPLE_LIGHT, PURPLE)

add_heading(doc, "Resumo do diagnóstico", 1)
add_table(doc, ["Sintoma", "Onde ocorre", "Leitura atual"], [
    ("Troca de aba demora", "GET de catálogo, parse e capas", "A aba já carrega em thread, mas o backend ainda pode responder lentamente; detalhes e progresso continuam síncronos."),
    ("HTTP 502 ao iniciar", "GET /api/play/:itemId ou gateway", "O backend devolve 502 quando resolveStreamUrl não consegue renovar nenhuma fonte; isso acontece antes do decoder."),
    ("Software fecha", "Abertura HLS/FFmpeg/libcurl/decoder", "Falha nativa não capturada. A diferença MP4 funcionando versus HLS fechando concentra a investigação no pipeline HLS e em memória/threads."),
    ("Animes funcionam", "Frequentemente embed resolvido para MP4", "Usam um caminho mais simples: uma URL/arquivo e HTTPS nativo do FFmpeg, sem várias renditions HLS simultâneas."),
], [2100, 2640, 4620])

add_callout(doc, "O que ainda não sabemos", "Sem a linha “Última etapa” do diagnóstico, o item/episódio exato e os logs correlacionados do backend, não é possível distinguir com certeza entre erro de fonte, falta de memória, falha do callback AVIO ou decoder. O plano abaixo foi desenhado para produzir essa evidência.", GOLD_LIGHT, GOLD)

doc.add_page_break()
add_heading(doc, "1. Visão geral do sistema", 1)
add_body(doc, "O aplicativo não recebe o catálogo e o vídeo na mesma requisição. O caminho é dividido em quatro blocos: descoberta do catálogo, detalhe da obra, criação da sessão e entrega dos bytes de mídia. Filmes e episódios chegam ao mesmo player, mas o formato e o resolvedor escolhidos pelo backend alteram completamente o transporte.")

add_table(doc, ["Etapa", "Cliente Switch", "Backend / mídia"], [
    ("1. Catálogo", "Escolhe a aba e faz GET autenticado.", "Monta hero e prateleiras consultando SQLite."),
    ("2. Detalhe", "Abre filme ou série e escolhe filme/episódio.", "Retorna sinopse, temporadas, episódios e progresso."),
    ("3. Sessão", "POST /api/stream/:itemId.", "Valida conta, fonte, plano e telas; cria session_id e play_url."),
    ("4. Gate de mídia", "FFmpeg/libcurl abre play_url.", "GET /api/play valida token, resolve URL real e responde 302."),
    ("5. Reprodução", "Abre MP4 ou HLS, decodifica e desenha/toca.", "R2/CDN entrega manifesto, segmentos, áudio e legendas."),
    ("6. Controle", "Heartbeat, progresso, refresh/fail e stop.", "Mantém sessão, histórico e troca de fonte."),
], [1400, 3580, 4380])

add_heading(doc, "1.1 Fluxo resumido", 2)
add_code(doc, "Controle -> enter_tab -> GET catálogo -> landing_apply -> open_item -> GET detalhe -> POST /api/stream/:id -> GET /api/play/:id -> HTTP 302 -> R2/CDN -> FFmpeg -> decoder -> SDL")
add_body(doc, "Autenticação: todas as chamadas de catálogo, sessão, progresso e heartbeat enviam o token da conta como Bearer. A URL /api/play usa um token curto na própria query; a URL final gerenciada do R2 recebe outro token, ligado ao prefixo da mídia, usuário, dispositivo e sessão.")

add_heading(doc, "2. Como cada aba é carregada", 1)
add_table(doc, ["Aba", "Requisição inicial", "Formato da resposta"], [
    ("Filmes", "GET /api/catalog/tab-home?tab=movie", "favoritos, prontos, hero, recent e shelves"),
    ("Séries", "GET /api/catalog/tab-home?tab=series", "favoritos, prontos, hero, recent e shelves"),
    ("Animes", "GET /api/catalog/anime-home", "continueWatching, favoritos, atualizados, populares, dublados, filmes e genreShelves"),
    ("Doramas", "GET /api/catalog/tab-home?tab=dorama", "favoritos, prontos, hero, recent e shelves por gênero"),
], [1500, 3480, 4380])

add_heading(doc, "2.1 Passo a passo no Switch", 2)
for text in [
    "L/R chama enter_tab(tab). A tela selecionada muda imediatamente.",
    "load_landing verifica g_land_cache[tab]. Se existir, reaplica o JSON sem rede.",
    "Sem cache, landing_start cria a thread catalog-fetch e mantém a interface desenhando.",
    "A thread chama api_get_timeout com 6 s para conexão e 30 s no total.",
    "net_request_timeout monta a URL base https://nplay.tonserverlocal.uk, adiciona Bearer, valida TLS pela CA embutida e recebe o corpo em memória.",
    "cJSON_Parse converte o corpo. pump_landing transfere o JSON para o cache da aba.",
    "landing_apply cria hero e rails. As capas são solicitadas separadamente e viram texturas aos poucos.",
]:
    add_number(doc, text)

add_callout(doc, "Risco de lentidão ainda presente", "A thread impede o congelamento do quadro, mas não reduz o tempo real do endpoint. Séries e Doramas executam várias consultas SQLite e várias prateleiras; Anime executa ainda mais agrupamentos. Sem cache HTTP/backend e métricas p95, o conteúdo pode continuar demorando para aparecer.", GOLD_LIGHT, GOLD)

add_heading(doc, "2.2 Pontos ainda síncronos", 2)
add_bullet(doc, "Filme: open_movie_details chama GET /api/catalog/movie/:id/info na thread da interface; o backend pode esperar até 7 s pelo enriquecimento.")
add_bullet(doc, "Série/Anime/Dorama: open_series chama GET /api/catalog/series/:id na thread da interface, com timeout total de 15 s.")
add_bullet(doc, "Antes do player: GET /api/sync/progress/:itemId também ocorre de forma síncrona.")
add_bullet(doc, "Busca, favoritos e algumas ações de lista ainda realizam rede na mesma thread que lê o controle.")
add_body(doc, "Isso explica por que a aba pode responder melhor, mas abrir uma obra, trocar uma versão de áudio ou iniciar o vídeo ainda aparenta travamento.")

doc.add_page_break()
add_heading(doc, "3. Classificação do item selecionado", 1)
add_body(doc, "Ao pressionar A, open_item decide se o cartão representa um filme, uma série ou um episódio. O campo kind tem prioridade; series_id também pode redirecionar diretamente para o detalhe de série.")
add_table(doc, ["Condição do cartão", "Ação"], [
    ("kind = live ou episode", "Chama resolve_and_play diretamente com o id do item."),
    ("series_id numérico", "Abre GET /api/catalog/series/:series_id."),
    ("kind = series ou rail marcada como série", "Abre GET /api/catalog/series/:id."),
    ("Demais casos", "Abre GET /api/catalog/movie/:id/info e mostra detalhe de filme."),
], [3300, 6060])

add_heading(doc, "3.1 Filme", 2)
add_body(doc, "O detalhe usa /api/catalog/movie/:id/info. Se info_at ainda não existe, o backend tenta enriquecer provedor/TMDB/TVDB e espera no máximo cerca de 7 s; depois responde com o que já existe. Ao selecionar Assistir, o id do próprio filme segue para resolve_and_play.")

add_heading(doc, "3.2 Série", 2)
add_body(doc, "O detalhe usa /api/catalog/series/:id. O backend monta versões de áudio, temporadas, episódios, metadados, progresso e imagens. O usuário seleciona um episódio; o player recebe o id de catalog_items daquele episódio, não o id da série.")

add_heading(doc, "3.3 Anime", 2)
add_body(doc, "Anime usa a mesma tela e o mesmo endpoint de detalhe de série. A diferença aparece na escolha da fonte: o backend prioriza mídia gerenciada R2; depois embeds; depois torrents e outras fontes. Um embed Hinatasoul pode ser resolvido para MP4 direto, que segue pelo caminho mais simples do player.")

add_heading(doc, "3.4 Dorama", 2)
add_body(doc, "Dorama também é série no cliente. Usa tab-home?tab=dorama, depois /series/:id e por fim o id do episódio em /stream/:itemId. Dependendo da fonte ativa, pode cair em HLS R2, embed resolvido ou outra fonte; portanto não é seguro assumir que todo dorama percorre o mesmo transporte.")

add_heading(doc, "4. Criação da sessão de reprodução", 1)
add_body(doc, "resolve_and_play desenha PREPARANDO e chama api_resolve_playback. Essa função envia POST /api/stream/:itemId com conexão máxima de 8 s e tempo total de 20 s.")
add_heading(doc, "4.1 O que o backend valida", 2)
for text in [
    "O item existe e, se for episódio, possui número de episódio válido.",
    "Há uma fonte ativa segundo a ordenação do backend.",
    "A conta tem acesso à fonte e não excedeu limite diário ou telas simultâneas.",
    "Cria ou reaproveita playback_sessions para o dispositivo.",
    "Classifica a entrega: r2 ou upstream; classifica container: m3u8, mp4, torrent ou embed.",
    "Devolve session_id, source_id e play_url. Para mídia comum, play_url é /api/play/:itemId com token curto.",
]:
    add_number(doc, text)

add_table(doc, ["Campo", "Uso no Switch", "Falha se incorreto"], [
    ("item_id", "Progresso, sessão e fallback", "Abre outra obra ou não encontra fonte"),
    ("session_id", "Heartbeat, refresh e limite de telas", "403/404 em renovação"),
    ("source_id", "Mantém ou troca a fonte correta", "502, 409 ou fallback incorreto"),
    ("delivery", "Ativa otimizações R2", "Sondagem excessiva ou caminho errado"),
    ("container", "Escolhe HLS versus arquivo", "MP4 tratado como HLS ou HLS como arquivo"),
    ("play_url", "Entrada real do FFmpeg", "Falha imediata ao abrir"),
], [1700, 3200, 4460])

add_heading(doc, "4.2 Seleção de fonte no backend", 2)
add_body(doc, "A ordem atual é: fonte gerenciada R2; embed; torrent classificado; demais fontes diretas. Dentro do conjunto, saúde, idioma, seeders, prioridade e fail_count participam do ranking. O cartão r2_ready é apenas um indicador de catálogo; o diagnóstico deve registrar a fonte realmente escolhida no POST /stream.")

add_callout(doc, "Possível origem do 502", "O POST /stream normalmente apenas cria a sessão. O player abre play_url logo depois. Em GET /api/play, resolveStreamUrl tenta descriptografar e renovar as fontes ativas; se todas lançarem erro, o backend responde HTTP 502 com “Não foi possível renovar a URL de stream desta fonte”. Isso acontece antes de FFmpeg receber o manifesto válido.", RED_LIGHT, RED)

add_heading(doc, "4.3 Trabalho duplicado no início", 2)
add_body(doc, "O handler POST /stream chama preloadStreamUrl em segundo plano e responde imediatamente. Logo depois, o Switch abre /api/play, que chama resolveStreamUrl novamente. Existe proteção inflightRefresh para fontes site-json renováveis, mas o desenho ainda produz duas tentativas próximas e duas trilhas de log. Em servidor lento ou fonte instável, isso aumenta ruído, carga e dificuldade de correlação.")

doc.add_page_break()
add_heading(doc, "5. Do play_url ao primeiro quadro", 1)
add_heading(doc, "5.1 Gate /api/play", 2)
for text in [
    "Valida token curto, item, usuário, dispositivo, sessão e source_id.",
    "resolveStreamUrl percorre fontes ativas e chama getFreshSiteStream quando a fonte exige renovação.",
    "Para mídia gerenciada, autoriza a URL final com token de longa duração.",
    "Responde HTTP 302 para R2/CDN ou para /api/hls quando a origem precisa de proxy.",
    "libcurl/FFmpeg segue o redirecionamento e começa a ler o recurso real.",
]:
    add_number(doc, text)

add_heading(doc, "5.2 Caminho MP4", 2)
add_body(doc, "MP4 remoto usa o protocolo HTTPS nativo do FFmpeg. Há uma única fonte principal; o demuxer MOV/MP4 usa Range/seek conforme necessário. É o caminho típico dos animes resolvidos de embed para MP4 e, segundo o relato, é o caminho que ainda consegue tocar.")
add_code(doc, "play_url -> 302 -> arquivo MP4 -> avformat_open_input -> avformat_find_stream_info -> decoder -> SDL")

add_heading(doc, "5.3 Caminho HLS", 2)
add_body(doc, "Filmes e muitos episódios usam container=m3u8. O demuxer HLS precisa abrir vários recursos: master, playlists filhas, init.mp4, segmentos de vídeo, segmentos de áudio e eventualmente legendas. O AVFormatContext instala io_open/io_close2; cada recurso HTTP(S) é entregue a nplay_curl_avio_open_hls.")
add_code(doc, "play_url -> 302 -> master.m3u8 -> child video + child audio (+ legendas) -> init.m4s -> segmentos .m4s -> demux -> decode")

add_table(doc, ["Componente HLS", "Comportamento atual"], [
    ("AVIO por recurso", "Cada URL recebe CurlIO, CURL easy, mutex, duas condições, thread produtora e buffers."),
    ("Buffer", "Ring de 2 MiB + bloco temporário de aproximadamente 320 KiB + AVIO de 64 KiB por recurso."),
    ("Primeiro byte", "cio_read pode aguardar aproximadamente 20 s no primeiro acesso."),
    ("Após iniciar", "Sem bytes, retorna EAGAIN a cada 300 ms para manter o player responsivo."),
    ("Timeout", "Conexão 20 s; low-speed 30 s; falhas transitórias podem tentar recuperar por até 120 s."),
    ("Sondagem", "2 MiB, 1,5 s, 64 pacotes; força H.264/AAC/WebVTT para R2."),
], [2600, 6760])

add_heading(doc, "5.4 Decoder e saída", 2)
add_body(doc, "Depois das faixas, o player escolhe vídeo, até 16 áudios e até 16 legendas de texto. Para H.264/HEVC tenta NVTEGRA; se falhar, reabre por CPU. O áudio é convertido para estéreo S16/48 kHz e enviado ao SDL. Só então heartbeat e progresso são liberados.")

add_heading(doc, "6. Recuperação e encerramento", 1)
add_table(doc, ["Evento", "Ação"], [
    ("Falha recuperável", "Códigos internos -2, -5 ou -10 entram no supervisor."),
    ("1ª recuperação", "POST /api/stream/session/:sessionId/refresh mantém a fonte."),
    ("2ª recuperação inicial", "POST /api/stream/session/:sessionId/fail tenta outra fonte."),
    ("Heartbeat", "A cada 20 s após pipeline_ready."),
    ("Progresso", "A cada 15 s e ao sair/seek."),
    ("Saída", "POST /api/stream/:itemId/stop encerra a sessão ativa."),
], [2500, 6860])

add_heading(doc, "7. Interpretação precisa dos erros", 1)
add_table(doc, ["Mensagem / código", "Origem provável", "Próxima verificação"], [
    ("Falha em Séries: HTTP 502", "GET tab-home atravessou gateway/backend e recebeu 502.", "Log da rota, tempo SQLite, memória/CPU e resposta do proxy."),
    ("Falha ao iniciar / HTTP 502", "GET /api/play não renovou nenhuma fonte ou o gateway encerrou a requisição.", "item_id, source_id, detalhe do 502 e logs resolveStreamUrl."),
    ("abrir playlist HLS: HTTP 502", "FFmpeg/libcurl abriu play_url e recebeu 502 antes de manifesto.", "Abrir a mesma URL autenticada no servidor; identificar fonte escolhida."),
    ("tempo esgotado", "avformat_open_input ou find_stream_info ultrapassou 20 s.", "Última etapa 03 ou 05, latência por recurso HLS."),
    ("software foi fechado", "Exceção nativa, corrupção, assert ou pressão de memória.", "Última etapa, crash report do Atmosphère e contadores de AVIO/memória."),
], [2350, 3550, 3460])

add_callout(doc, "Separação essencial", "HTTP 502 é uma resposta controlada de rede/backend. “O software foi fechado porque ocorreu um erro” é um encerramento nativo do processo. Mesmo que ambos apareçam ao pressionar Assistir, devem ter incidentes, logs e critérios de correção separados.", PURPLE_LIGHT, PURPLE)

doc.add_page_break()
add_heading(doc, "8. Marcadores persistentes do player", 1)
add_body(doc, "A versão 0.9.5 grava a última etapa em sdmc:/switch/.nplay-player-boot.txt. A mesma informação aparece em Configurações > X Diagnóstico após reiniciar o aplicativo.")
add_table(doc, ["Etapa", "Significado", "Se o software fechou aqui"], [
    ("01 início do player", "Entrou em player_play_internal.", "Falha muito precoce ou escrita/estado global."),
    ("02 contexto HLS libcurl pronto", "AVFormatContext criado e callbacks instalados.", "Falha antes de abrir URL."),
    ("03 abrindo fonte", "Dentro de avformat_open_input.", "Gate /api/play, redirect, master HLS ou callback AVIO."),
    ("04 fonte aberta", "Master/demuxer inicial aceito.", "Falha entre abertura e inspeção de faixas."),
    ("05 lendo faixas", "Dentro de avformat_find_stream_info.", "Renditions, segmentos init, múltiplos AVIO ou memória."),
    ("06 faixas prontas", "Stream info concluído.", "Enumeração ou seleção de áudio/legenda/vídeo."),
    ("07 decoder de vídeo pronto", "NVTEGRA ou CPU abriu codec.", "Configuração de áudio ou alocação seguinte."),
    ("08 áudio pronto", "Áudio tentado/configurado.", "Textura, frames, packets ou início do loop."),
    ("09 reproduzindo", "Loop principal iniciado.", "Decode de pacote/frame, upload NV12/IYUV, áudio ou thread concorrente."),
], [1600, 3000, 4760])

add_heading(doc, "9. Hipóteses priorizadas", 1)
add_heading(doc, "P0 - Caminho HLS customizado", 2)
add_bullet(doc, "Evidência: MP4 de anime funciona, enquanto filmes/séries HLS fecham.")
add_bullet(doc, "Risco: callbacks io_open/io_close2, múltiplas threads CurlIO e ciclo de vida controlado pelo demuxer.")
add_bullet(doc, "Teste decisivo: duas builds A/B do mesmo item R2 - uma com HTTP nativo FFmpeg, outra com AVIO libcurl - e registrar a etapa final.")

add_heading(doc, "P0 - Pressão de memória e quantidade de recursos", 2)
add_bullet(doc, "Cada recurso HLS pode consumir mais de 2,3 MiB antes de estruturas libcurl/FFmpeg.")
add_bullet(doc, "Um master com vídeo, vários áudios e legendas pode manter vários recursos simultâneos.")
add_bullet(doc, "O app ainda possui cache de até 160 texturas de capa, decoder, frames, áudio e JSONs de cinco abas.")
add_bullet(doc, "Teste decisivo: registrar memória livre e número/bytes de CurlIO antes de cada etapa; repetir com cache de capas esvaziado e HLS de uma única rendition.")

add_heading(doc, "P0 - Fonte marcada como pronta, mas não reproduzível", 2)
add_bullet(doc, "r2_ready no catálogo não prova que a fonte escolhida, o manifesto e todos os segmentos continuam válidos.")
add_bullet(doc, "O POST /stream deve registrar delivery, container e source_id efetivos, sem registrar token/URL assinada.")
add_bullet(doc, "Teste decisivo: validar master, child, init e primeiro segmento do source_id escolhido imediatamente antes do teste no console.")

add_heading(doc, "P1 - 502 no resolvedor", 2)
add_bullet(doc, "resolveStreamUrl devolve 502 quando todas as fontes geram exceção.")
add_bullet(doc, "preload e abertura real acontecem quase juntos; logs precisam de request_id para separar tentativas.")
add_bullet(doc, "Teste decisivo: medir POST /stream e GET /play separadamente; um R2 saudável deve responder 302 rapidamente e sem resolver site externo.")

add_heading(doc, "P1 - Operações síncronas na interface", 2)
add_bullet(doc, "Detalhes de filme/série e progresso ainda podem bloquear de 7 a 15 s.")
add_bullet(doc, "Teste decisivo: cronômetro por ação e migração dessas três chamadas para estado assíncrono cancelável.")

add_heading(doc, "P1 - Catálogo pesado no backend", 2)
add_bullet(doc, "tab-home executa consultas por prateleira; anime-home monta muitos conjuntos e enriquecimentos locais.")
add_bullet(doc, "Teste decisivo: medir tempo SQL, serialização, tamanho do JSON e tempo total p50/p95 por aba, a frio e quente.")

doc.add_page_break()
add_heading(doc, "10. Instrumentação mínima recomendada", 1)
add_body(doc, "O próximo build deve produzir um registro JSONL local e logs correlacionados no backend. Nunca registrar token, play_url completo, senha ou URL assinada.")
add_heading(doc, "10.1 Evento no Switch", 2)
add_code(doc, '{"at":"...","event":"hls_io_open","item_id":123,"session_id":45,"source_id":67,"delivery":"r2","container":"m3u8","stage":"05","active_avio":3,"avio_bytes":7077888,"free_memory":123456789,"http_status":302,"ffmpeg_rc":0}')
add_bullet(doc, "Gerar correlation_id no início de resolve_and_play e enviá-lo como cabeçalho X-Nplay-Trace nas chamadas de controle.")
add_bullet(doc, "Contar AVIO abertos/fechados, bytes reservados e pico simultâneo.")
add_bullet(doc, "Registrar somente host e tipo do recurso: play gate, master, child, init, segment, subtitle.")
add_bullet(doc, "Registrar avformat_open_input, find_stream_info, abertura de decoder, primeira leitura e primeiro quadro.")
add_bullet(doc, "Registrar memória livre antes/depois de cache, abertura HLS e decoder.")

add_heading(doc, "10.2 Evento no backend", 2)
add_code(doc, '{"trace":"...","route":"GET /api/play/:itemId","item_id":123,"source_id":67,"delivery":"r2","resolver_ms":18,"result":"302","target_kind":"r2"}')
add_bullet(doc, "Tempo total e status de tab-home, anime-home, series/:id, stream/:id e play/:id.")
add_bullet(doc, "Para 502: classe do erro e source_id tentados; remover URL, query e segredo.")
add_bullet(doc, "Tempo de SQLite separado do tempo de renovação externa.")
add_bullet(doc, "Contador de chamadas preload versus chamadas reais para o mesmo item/source.")

add_heading(doc, "11. Plano de investigação passo a passo", 1)
add_heading(doc, "Fase A - Reproduzir e classificar", 2)
for text in [
    "Escolher quatro amostras fixas: um anime MP4 conhecido, um filme R2 HLS, um episódio de série R2 HLS e um dorama; registrar ids.",
    "Reiniciar o Switch antes de cada amostra e anotar versão 0.9.5, modo de execução e condição da rede.",
    "Executar uma vez e fotografar Configurações > X Diagnóstico após qualquer fechamento.",
    "Separar o resultado em catálogo, detalhe, POST /stream, GET /play, abertura HLS e primeiro quadro.",
]:
    add_number(doc, text)

add_heading(doc, "Fase B - Confirmar servidor e fonte", 2)
for text in [
    "No mesmo minuto, localizar o item_id/source_id nos logs do backend.",
    "Confirmar se POST /stream retornou 200 e qual descriptor foi escolhido.",
    "Confirmar se GET /play retornou 302, 404, 403 ou 502 e quanto demorou.",
    "Para R2, testar master, uma child, init e primeiro segmento sem imprimir o token.",
    "Se /play for 502, corrigir fonte/resolvedor antes de tocar no decoder do Switch.",
]:
    add_number(doc, text)

add_heading(doc, "Fase C - Isolar o fechamento HLS", 2)
for text in [
    "Instrumentar active_avio, bytes reservados e memória livre.",
    "Testar master simplificado com uma faixa de vídeo e uma de áudio, sem legendas.",
    "Comparar build A com HTTP nativo FFmpeg e build B com callbacks libcurl, usando o mesmo manifesto.",
    "Desabilitar temporariamente NVTEGRA somente no build de diagnóstico; comparar etapa 06/07/09.",
    "Repetir com cache de capas limpo/reduzido. Se a falha mudar, priorizar orçamento de memória.",
]:
    add_number(doc, text)

add_heading(doc, "Fase D - Remover latência percebida", 2)
for text in [
    "Mover detalhe de filme, detalhe de série e progresso para jobs assíncronos canceláveis.",
    "Adicionar cache backend curto para landings e ETag/versão do catálogo.",
    "Evitar enriquecimento externo no caminho crítico de abrir uma obra.",
    "Pré-carregar apenas vizinhos e cancelar downloads de capa da aba abandonada.",
]:
    add_number(doc, text)

add_heading(doc, "12. Matriz de teste", 1)
add_table(doc, ["Caso", "Pré-condição", "Evidência", "Resultado esperado"], [
    ("Filme R2", "source_id R2 validado", "trace + etapa + backend", "302 rápido; primeiro quadro; 30 s estáveis"),
    ("Série R2", "episódio real > 0", "mesmas evidências", "áudio/vídeo; autoplay não testado ainda"),
    ("Anime MP4", "embed resolve para mp4", "trace container=mp4", "baseline funcional"),
    ("Dorama", "anotar tipo real", "descriptor", "comportamento coerente com mp4/hls"),
    ("Fonte inválida", "primeira fonte falha", "refresh/fail", "fallback controlado, sem fechar"),
    ("Rede lenta", "latência/perda controlada", "buffering + B", "UI responde e permite cancelar"),
    ("Aba fria", "cache vazio", "p50/p95 + bytes", "input imediato; conteúdo aparece <= 3 s alvo"),
    ("Aba quente", "cache preenchido", "tempo local", "conteúdo imediato, sem nova GET"),
], [1500, 2650, 2250, 2960])

add_heading(doc, "13. Critérios de aceite", 1)
add_bullet(doc, "Nenhum fechamento nativo em 20 aberturas alternadas de filme/série/anime/dorama.")
add_bullet(doc, "Filme e episódio HLS reproduzem 30 minutos, com áudio, três seeks e retorno por B.")
add_bullet(doc, "Troca visual de aba responde em até 250 ms; carregamento de rede nunca bloqueia controles.")
add_bullet(doc, "Landing fria p95 até 3 s na rede de referência; detalhe local p95 até 2 s.")
add_bullet(doc, "POST /stream R2 p95 até 1 s e GET /play R2 p95 até 500 ms.")
add_bullet(doc, "Toda falha exibe etapa, item_id e classe do erro sem expor token ou URL assinada.")
add_bullet(doc, "HTTP 502 de fonte não derruba o processo; tenta fallback ou retorna mensagem controlada.")

add_heading(doc, "14. Ordem recomendada de trabalho", 1)
add_table(doc, ["Prioridade", "Entrega", "Motivo"], [
    ("1", "Instrumentação + correlação", "Sem isso cada mudança continua sendo hipótese."),
    ("2", "Separar 502 de crash nativo", "Evita alterar FFmpeg para um problema de fonte ou vice-versa."),
    ("3", "A/B do transporte HLS", "Confirma ou elimina o principal diferencial frente ao MP4."),
    ("4", "Orçamento de memória", "Fecha a hipótese de múltiplos buffers, capas e decoder."),
    ("5", "Assincronizar detalhes/progresso", "Remove congelamentos ainda objetivos no cliente."),
    ("6", "Otimizar/cachear landings", "Reduz o tempo real de catálogo depois de estabilizar o fluxo."),
], [1200, 3440, 4720])

add_callout(doc, "Recomendação", "Não publicar outra correção ampla do player antes de obter, para uma única reprodução que fecha: item_id, source_id, delivery, container, status de /stream, status de /play, última etapa, pico de AVIO e memória livre. Esse conjunto transforma o próximo ajuste em uma correção verificável.", GREEN_LIGHT, GREEN)

doc.add_page_break()
add_heading(doc, "Apêndice A - Arquivos e pontos de entrada", 1)
add_table(doc, ["Arquivo", "Responsabilidade / ponto atual"], [
    ("source/main.c", "enter_tab ~2017; catálogo ~422-557; resolve_and_play ~683; open_series ~809; open_item ~817."),
    ("source/screen_movie.c", "fetch_movie_details ~61; open_movie_details ~91; Assistir chama resolve_and_play ~312."),
    ("source/api.c", "POST /api/stream ~111-168; refresh/fail ~169-229; heartbeat/progresso/stop ~230-266."),
    ("source/net.c", "TLS, Bearer, timeouts, CURLSH e download HTTP."),
    ("source/player.c", "HLS io_open ~99; abertura/decoder ~576; supervisor player_run ~1387."),
    ("source/curl_avio.c", "producer ~124; cio_read ~185; alocação do perfil HLS ~249."),
    ("C:/iptv/src/routes/catalog.js", "tab-home ~390; anime-home ~501; series/:id ~678; movie/:id/info ~921."),
    ("C:/iptv/src/routes/stream.js", "seleção de fontes ~117; resolveStreamUrl ~177; POST /stream ~374; GET /play ~634; refresh/fail ~879."),
], [3300, 6060])

add_heading(doc, "Apêndice B - Ficha de coleta", 1)
for label in [
    "Versão/commit do NRO:", "Data e hora do console:", "Modo de abertura (álbum ou title override):",
    "Aba e título:", "item_id / series_id / episódio:", "Mensagem exata:", "Última etapa:",
    "POST /stream status e duração:", "GET /play status e duração:", "delivery / container / source_id:",
    "Pico active_avio / bytes / memória livre:", "Crash report do Atmosphère:", "Observações:",
]:
    p = doc.add_paragraph()
    set_para(p, before=2, after=8, line=1.0)
    set_run(p.add_run(label + "  "), bold=True, color=NAVY)
    set_run(p.add_run("____________________________________________________________"), color="AAB2BD")

add_heading(doc, "Apêndice C - Limites desta análise", 1)
add_body(doc, "Este documento descreve o código local no commit 9ca16fd e o backend disponível em C:/iptv em 31/08/2026. Não substitui teste no Nintendo Switch real. Não foram usados tokens da conta nem URLs assinadas em testes externos durante esta rodada. Linhas podem mudar após novos commits; nomes de funções e endpoints são a referência mais estável.")

# Evita linhas viúvas e padroniza fontes de todos os runs criados por estilos.
for p in doc.paragraphs:
    if p.style and p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True

doc.core_properties.title = "Nplay Switch - Fluxo de reprodução e plano de investigação"
doc.core_properties.subject = "Catálogo, sessão, HLS, HTTP 502 e fechamento nativo"
doc.core_properties.author = "Nplay"
doc.core_properties.keywords = "Nintendo Switch, Nplay, FFmpeg, HLS, libcurl, diagnóstico"
doc.save(OUT)
print(OUT)
